# encoding=utf-8
import docx
import collections
import jieba
import re
import string
import xlwt
import os

def CountWordsNum(file):
    result = [0, 0, 0, 0, 0, 0, 0, 0]
    dict = ['相关公众','公众','相关消费者','普通消费者','消费者','普通购买者','购买者','他人']
    PATTEN = ["[\u4e00-\u9fa5]?(相关公众)", 
          "[\u4e00-\u9fa5]?(公众)",
          "[\u4e00-\u9fa5]?(相关消费者)",
          "[\u4e00-\u9fa5]?(普通消费者)",
          "[\u4e00-\u9fa5]?(消费者)",
          "[\u4e00-\u9fa5]?(普通购买者)",
          "[\u4e00-\u9fa5]?(购买者)",
          "[\u4e00-\u9fa5]?(他人)"
         ]
    for p in file.paragraphs:
    #words = jieba.cut(p.text, cut_all = True)
    #for w in words:
    #    if w == "相关公众":
    #        print("once!")
        i = 0
        for patten in PATTEN:
            pa = re.compile(patten)
            m = pa.findall(p.text)
            if len(m) != 0:
                #print(m)
                result[i] += len(m)
            i += 1
    result[1] = result[1] - result[0]
    result[4] = result[4] - result[3] - result[2]
    result[6] = result[6] - result[5]
    
    maxnum = max(result)
    resstr = ""
    j = 0
    for r in result:
        if r == maxnum:
           resstr = resstr + dict[j] + " "
        j += 1
    return resstr

def JudgeCorporation(file):
    k = 0
    for p in file.paragraphs:
        if len(p.text) != 0:
            break
        k += 1
    resstr = []
    title = file.paragraphs[k]
    title = title.text
    subtitle = title.split("诉")
    PATTEN = "[\u4e00-\u9fa5]?(公司)"
    patten = re.compile(PATTEN)

    res1 = patten.search(subtitle[0])
    if res1 == None:
        print("原告自然人")
        resstr.append("原告自然人")
    else:
        print("原告法人")
        resstr.append("原告法人")

    res2 = patten.search(subtitle[1])
    if res2 == None:
        print("被告自然人")
        resstr.append("被告自然人")
    else:
        print("被告法人")
        resstr.append("被告法人")
    return resstr


def CharacterofCase(file):
    start = 0
    end = 0
    j = 0

    for p in file.paragraphs:
        #去除空格
        s = p.text.split()
        if len(s) != 0:
            if s[0] == "【裁判理由】":
                start = j
            if s[0] == "本案法律依据":
                end = j
                break
        j += 1

    PATTEN1 = "[\u4e00-\u9fa5]?(混淆)"
    PATTEN2 = "[\u4e00-\u9fa5]?(近似)"
    PATTEN3 = "[\u4e00-\u9fa5]?(。)"
    patten1 = re.compile(PATTEN1)
    patten2 = re.compile(PATTEN2)
    patteb3 = re.compile(PATTEN3)
    flag1 = 0
    flag2 = 0

    for n in range(start + 1, end):
        p = file.paragraphs[n]
        res1 = patten1.search(p.text)
        res2 = patten2.search(p.text)
        if res1 != None:
            flag1 = 1
        if res2 != None:
            flag2 = 1

    can_judge = 0
    if flag1 == 1 and flag2 == 0:
        print("混淆标准")
        return("混淆标准")
    elif flag1 == 0 and flag2 == 1:
        print("近似标准")
        return("近似标准")
    elif flag1 == 0 and flag2 == 0:
        print("无关案件")
        return("无关案件")
    else:
        for n in range(start + 1, end):
            p = file.paragraphs[n].text
            subp = p.split("。")
            for sp in subp:
                res_hx = patten1.search(sp)
                res_js = patten2.search(sp)
                if res_hx != None and res_js != None:
                    doublesub = sp.split("混淆")
                    res_former = patten2.search(doublesub[0])
                    if res_former == None:
                        print("混淆导致相似")
                        can_judge = 1
                        return("混淆导致相似")
                    else: 
                        print("相似导致混淆")
                        can_judge = 1
                        return("相似导致混淆")
 
        if can_judge == 0:
            print("无法判断")
            return("无法判断")

def set_style(name, height, bold=False):
    style = xlwt.XFStyle()   # 初始化样式
    font = xlwt.Font()       # 为样式创建字体
    font.name = name
    font.bold = bold
    font.color_index = 4
    font.height = height

    style.font = font
    return style


def write_excel(path, rows):
    #创建工作表
    workbook = xlwt.Workbook(encoding='utf-8')
    #创建sheet
    sheet = workbook.add_sheet('sheet1')
    firstline = ["原告性质", "被告性质", "最高频词汇", "案件性质"]
    for col in range(len(firstline)):
        sheet.write(0, col, firstline[col], set_style('黑体', 220, True))
    rownum = 1
    for r in rows:
        col = 0
        for rr in r:
            sheet.write(rownum, col, rr)
            col += 1
        rownum += 1
    workbook.save(path)

def ScanFile(filepath):
    allfiles = os.listdir(filepath)
    return allfiles

if __name__ == '__main__':
    #Open file .docx
    #file = docx.Document('/Users/eagleying/Downloads/Cases/指导案例82号 王碎永诉深圳歌力思服饰股份有限公司、杭州银泰世纪百货有限公司侵害商标权纠纷案.docx')
    #print("段落数:"+str(len(file.paragraphs)))
    print("请输入您word案例所在的路径,例如:")
    print("Mac用户: /Users/eagleying/Downloads/Cases/")
    print("Windows用户: D:\\Case\\")
    filepath = input()
    print("请输入Excel文件名,如result.xls")
    filename = input()
    print("请输入生成Excel文件的路径:")
    excelpath = input()
    #filepath = '/Users/eagleying/Downloads/Cases/'
    allfiles = ScanFile(filepath)
    rows = []
    for f in allfiles:
        file = docx.Document(filepath + f)
        row = []
        for r in JudgeCorporation(file):
            row.append(r)
        row.append(CountWordsNum(file))
        row.append(CharacterofCase(file))
        rows.append(row)
    #path = '/Users/eagleying/Downloads/resuults.xls'
    path = excelpath + filename
    write_excel(path, rows)
    print("生成成功^o^")    
